import streamlit as st
import google.generativeai as genai
from docx import Document
import os
import time
import re

# --- НАСТРОЙКИ СТРАНИЦЫ ---
st.set_page_config(page_title="AI Следователь", page_icon="🕵️‍♂️", layout="wide")

# --- БЛОК БЕЗОПАСНОСТИ (API KEY) ---
try:
    # 1. Сначала пытаемся найти ключ в настройках сервера (для облака)
    API_KEY = st.secrets["GOOGLE_API_KEY"]
except:
    # 2. Если не нашли (ты на своем ПК), используем этот ключ:
    # !!! ВСТАВЬ СВОЙ КЛЮЧ ВНУТРЬ КАВЫЧЕК НИЖЕ !!!
    API_KEY = "ВСТАВЬ_СЮДА_ТВОЙ_КЛЮЧ" 

# --- СОСТОЯНИЕ (ПАМЯТЬ) ---
if 'case_files' not in st.session_state:
    st.session_state['case_files'] = [] 
if 'brief_text' not in st.session_state:
    st.session_state['brief_text'] = ""
if 'generated_filename' not in st.session_state:
    st.session_state['generated_filename'] = "Case_Brief.docx"

# --- ФУНКЦИИ ---

def sanitize_filename(name):
    """Чистит имя файла от мусора"""
    clean_name = re.sub(r'[\\/*?:"<>|]', "", name).strip()
    if len(clean_name) < 3: clean_name = f"Case_{int(time.time())}"
    if not clean_name.endswith(".docx"): clean_name += ".docx"
    return clean_name

def analyze_case_files(file_paths):
    """Анализирует список файлов"""
    genai.configure(api_key=API_KEY)
    status = st.status("🕵️‍♂️ Изучаю материалы дела (все записи)...", expanded=True)
    
    try:
        uploaded_files_for_gemini = []
        
        # 1. Загрузка файлов
        for path in file_paths:
            status.write(f"📤 Читаю файл: {os.path.basename(path)}...")
            upl_file = genai.upload_file(path=path)
            
            while upl_file.state.name == "PROCESSING":
                time.sleep(1)
                upl_file = genai.get_file(upl_file.name)
            
            uploaded_files_for_gemini.append(upl_file)

        # 2. Анализ
        status.write("🧠 Сопоставляю факты и формулиру вопросы...")
        model = genai.GenerativeModel('models/gemini-2.0-flash')
        
        prompt = """
        Ты - профессиональный юрист-следователь.
        Тебе передали аудиозаписи по ОДНОМУ делу.
        
        Твоя задача:
        1. Объединить информацию из ВСЕХ файлов.
        2. Придумай имя файла (FILENAME) на латинице.
        3. Составь Сводку и список вопросов.

        СТРУКТУРА ОТВЕТА:
        FILENAME: [Name_on_Latin.docx]

        ЮРИДИЧЕСКАЯ СВОДКА (ОБЪЕДИНЕННАЯ)
        
        1. 📄 ОБСТОЯТЕЛЬСТВА:
           [Связный рассказ]

        2. 🔑 КЛЮЧЕВЫЕ ФАКТЫ:
           - Стороны:
           - Даты:
           - Суммы:
        
        3. ❓ ВОПРОСЫ КЛИЕНТУ (ЧТО СПРОСИТЬ):
           [Список вопросов для уточнения]

        4. 🚀 ПЛАН ДЕЙСТВИЙ:
           [Рекомендации]
        """
        
        request_content = [prompt] + uploaded_files_for_gemini
        response = model.generate_content(request_content)
        text_full = response.text
        
        # Парсинг имени файла
        lines = text_full.split('\n')
        suggested_name = "Case_Brief.docx"
        content_start = 0
        if lines and lines[0].startswith("FILENAME:"):
            suggested_name = sanitize_filename(lines[0].replace("FILENAME:", ""))
            content_start = 1
            
        clean_text = "\n".join(lines[content_start:]).strip()
        
        status.update(label="✅ Дело обновлено!", state="complete", expanded=False)
        return clean_text, suggested_name
        
    except Exception as e:
        status.error(f"Ошибка: {e}")
        return None, None

def create_doc(text, filename):
    doc = Document()
    doc.add_heading('МАТЕРИАЛЫ ДЕЛА', 0)
    doc.add_paragraph(text)
    doc.save(filename)

# --- ИНТЕРФЕЙС ---

st.title("🕵️‍♂️ AI Следователь: Сбор улик")

col1, col2 = st.columns([1, 1])

# ЛЕВАЯ КОЛОНКА
with col1:
    st.subheader("1. Добавить улики")
    
    tab_mic, tab_up = st.tabs(["🎙️ Записать", "📂 Загрузить"])
    new_audio_source = None
    
    with tab_mic:
        mic_val = st.audio_input("Микрофон", key="mic_input")
        if mic_val: new_audio_source = mic_val
            
    with tab_up:
        file_val = st.file_uploader("Файл", type=["mp3", "wav", "ogg", "m4a", "opus"], key="file_input")
        if file_val: new_audio_source = file_val

    if st.button("➕ Добавить в дело и Обновить", type="primary"):
        if new_audio_source:
            ext = "wav"
            if hasattr(new_audio_source, "name"):
                ext = new_audio_source.name.split(".")[-1]
            
            filename = f"evidence_{int(time.time())}.{ext}"
            with open(filename, "wb") as f:
                f.write(new_audio_source.getbuffer())
            
            st.session_state['case_files'].append(filename)
            
            res_text, res_name = analyze_case_files(st.session_state['case_files'])
            
            if res_text:
                st.session_state['brief_text'] = res_text
                st.session_state['generated_filename'] = res_name
                st.rerun()
        else:
            st.warning("Сначала запишите или выберите файл!")

    st.divider()
    st.write(f"📎 **Файлов в деле: {len(st.session_state['case_files'])}**")
    
    if st.button("🗑️ Начать новое дело (Сброс)"):
        for f in st.session_state['case_files']:
            if os.path.exists(f): os.remove(f)
        st.session_state['case_files'] = []
        st.session_state['brief_text'] = ""
        st.rerun()

# ПРАВАЯ КОЛОНКА
with col2:
    st.subheader("2. Результат")
    
    if st.session_state['brief_text']:
        st.info(f"📁 Файл: **{st.session_state['generated_filename']}**")
        
        edited_text = st.text_area("Сводка", value=st.session_state['brief_text'], height=600)
        st.session_state['brief_text'] = edited_text
        
        if st.button("💾 Скачать (.docx)"):
            fname = st.session_state['generated_filename']
            create_doc(edited_text, fname)
            with open(fname, "rb") as f:
                st.download_button("📥 Скачать", f, file_name=fname)
    else:
        st.info("Нет данных. Добавьте первую запись.")