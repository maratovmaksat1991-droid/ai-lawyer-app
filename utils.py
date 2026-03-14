import os
import time
import io
import streamlit as st
import google.generativeai as genai
from docx import Document
import PyPDF2
from dotenv import load_dotenv

load_dotenv()


def get_api_key():
    """Получает API-ключ из .env или st.secrets."""
    key = os.getenv("GOOGLE_API_KEY")
    if not key:
        try:
            key = st.secrets["GOOGLE_API_KEY"]
        except Exception:
            key = ""
    return key


def configure_genai():
    """Настраивает Gemini API."""
    key = get_api_key()
    if not key:
        st.error("API-ключ не найден. Создайте файл .env с GOOGLE_API_KEY=ваш_ключ")
        st.stop()
    genai.configure(api_key=key)


def get_model():
    """Возвращает модель Gemini."""
    configure_genai()
    return genai.GenerativeModel('models/gemini-2.5-flash')


def extract_text_from_file(uploaded_file):
    """Извлекает текст из PDF, DOCX или TXT файла."""
    if uploaded_file is None:
        return ""

    text = ""
    try:
        name = uploaded_file.name.lower()
        if name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif name.endswith('.docx'):
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif name.endswith('.txt'):
            text = uploaded_file.read().decode("utf-8")

        if not text.strip():
            return "Файл пуст или не удалось извлечь текст."
    except Exception as e:
        return f"Ошибка при чтении файла: {e}"
    return text


def create_docx(content, title="Документ"):
    """Создаёт DOCX-файл в памяти и возвращает буфер."""
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def transcribe_audio(audio_bytes):
    """Транскрибирует аудио через Gemini."""
    model = get_model()
    temp_name = f"voice_{int(time.time())}.wav"
    try:
        with open(temp_name, "wb") as f:
            f.write(audio_bytes.getbuffer())
        upl = genai.upload_file(temp_name)
        while upl.state.name == "PROCESSING":
            time.sleep(0.5)
        res = model.generate_content(["Транскрибируй текст аудио.", upl])
        return res.text
    except Exception as e:
        st.error(f"Ошибка при транскрибации: {e}")
        return ""
    finally:
        if os.path.exists(temp_name):
            os.remove(temp_name)


def upload_audio_to_gemini(file_path):
    """Загружает аудиофайл в Gemini и возвращает объект файла."""
    upl = genai.upload_file(file_path)
    while upl.state.name == "PROCESSING":
        time.sleep(0.5)
    return upl


def cleanup_temp_files(file_paths):
    """Удаляет временные файлы."""
    for f in file_paths:
        if os.path.exists(f):
            os.remove(f)
