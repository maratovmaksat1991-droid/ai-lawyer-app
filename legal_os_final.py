import streamlit as st
import google.generativeai as genai
from docx import Document
import PyPDF2
import os
import time
import io

# --- 1. НАСТРОЙКИ СТРАНИЦЫ И ДИЗАЙН ---
st.set_page_config(page_title="Legal OS Pro", page_icon="⚖️", layout="wide")

# Внедряем CSS для красоты (Крупные кнопки, отступы, шрифты)
st.markdown("""
<style>
    /* Увеличиваем заголовки */
    h1 {
        color: #1E3A8A;
        font-family: 'Helvetica Neue', sans-serif;
    }
    h2, h3 {
        color: #374151;
    }
    /* Стиль для главных кнопок (синие) */
    div.stButton > button:first-child {
        background-color: #2563EB;
        color: white;
        border-radius: 10px;
        height: 50px;
        font-size: 18px;
        font-weight: bold;
        border: none;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        transition: all 0.3s;
    }
    div.stButton > button:first-child:hover {
        background-color: #1D4ED8;
        transform: translateY(-2px);
    }
    /* Стиль для контейнеров (карточки) */
    .stExpander {
        border: 1px solid #E5E7EB;
        border-radius: 8px;
    }
    /* Убираем лишние отступы сверху */
    .block-container {
        padding-top: 2rem;
    }
</style>
""", unsafe_allow_html=True)

# !!! КЛЮЧ API !!!
try:
    API_KEY = st.secrets["GOOGLE_API_KEY"]
except:
    API_KEY = "" 

genai.configure(api_key=API_KEY)

# --- ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ (БЕЗ ИЗМЕНЕНИЙ) ---
def extract_text_from_file(uploaded_file):
    text = ""
    try:
        if uploaded_file.name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages: text += page.extract_text() + "\n"
        elif uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            for para in doc.paragraphs: text += para.text + "\n"
        elif uploaded_file.name.endswith('.txt'):
            text = uploaded_file.read().decode("utf-8")
    except Exception as e: return f"Ошибка: {e}"
    return text

def create_docx(content, title="Документ"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def transcribe_audio(audio_bytes):
    model = genai.GenerativeModel('models/gemini-2.0-flash')
    temp_name = f"voice_{int(time.time())}.wav"
    with open(temp_name, "wb") as f: f.write(audio_bytes.getbuffer())
    upl = genai.upload_file(temp_name)
    while upl.state.name == "PROCESSING": time.sleep(0.5)
    res = model.generate_content(["Текст аудио.", upl])
    if os.path.exists(temp_name): os.remove(temp_name)
    return res.text

# --- МОДУЛЬ 1: АНАЛИЗАТОР (ДИЗАЙН) ---
def render_audio_analyzer():
    st.markdown("## 🕵️‍♂️ Анализатор Дела")
    st.caption("Режим «Следователь»: Соберите улики, и ИИ составит полную картину.")

    if 'case_files' not in st.session_state: st.session_state.case_files = []
    if 'brief_text' not in st.session_state: st.session_state.brief_text = ""

    col_left, col_right = st.columns([1, 1.2], gap="large")

    with col_left:
        with st.container(border=True):
            st.markdown("### 1. Добавить материалы")
            st.info("Можно добавлять записи по очереди. Сводка будет обновляться.", icon="ℹ️")
            
            tab_mic, tab_file = st.tabs(["🎙️ Записать голос", "📂 Загрузить файл"])
            new_source = None
            
            with tab_mic:
                mic_val = st.audio_input("Нажмите для записи", key="an_mic") 
                if mic_val: new_source = mic_val
            with tab_file:
                file_val = st.file_uploader("Выберите аудио (MP3, WA, OGG)", type=["mp3", "wav", "ogg", "m4a"])
                if file_val: new_source = file_val
            
            st.markdown("---")
            
            if st.button("➕ Добавить в дело", use_container_width=True, help="Нажмите, чтобы отправить аудио на анализ"):
                if new_source:
                    ext = new_source.name.split('.')[-1] if hasattr(new_source, 'name') else 'wav'
                    fname = f"evid_{int(time.time())}.{ext}"
                    with open(fname, "wb") as f: f.write(new_source.getbuffer())
                    
                    st.session_state.case_files.append(fname)
                    st.toast("✅ Файл добавлен к делу!")
                    
                    with st.spinner("🔍 ИИ анализирует новые факты..."):
                        files_gemini = []
                        for p in st.session_state.case_files:
                            uf = genai.upload_file(p)
                            while uf.state.name == "PROCESSING": time.sleep(0.5)
                            files_gemini.append(uf)
                        
                        model = genai.GenerativeModel('models/gemini-2.0-flash')
                        prompt = """
                        Роль: Юрист РК.
                        Задача: Сводный отчет по всем записям.
                        Структура:
                        1. 📋 СУТЬ СИТУАЦИИ.
                        2. 🔢 ФАКТЫ (Даты, Суммы, Имена).
                        3. ⚖️ ПРАВОВАЯ ОЦЕНКА (ГК/ГПК РК).
                        4. ❓ ЧТО УТОЧНИТЬ У КЛИЕНТА.
                        5. 🚀 ПЛАН ДЕЙСТВИЙ.
                        """
                        res = model.generate_content([prompt] + files_gemini)
                        st.session_state.brief_text = res.text
                else:
                    st.warning("Сначала выберите файл или запишите голос.")

            if st.session_state.case_files:
                st.write(f"📎 В деле файлов: **{len(st.session_state.case_files)}**")
                if st.button("🗑️ Сбросить всё", type="secondary"):
                    st.session_state.case_files = []
                    st.session_state.brief_text = ""
                    st.rerun()

    with col_right:
        with st.container(border=True):
            st.markdown("### 2. Сводка по делу")
            if st.session_state.brief_text:
                st.markdown(st.session_state.brief_text)
                st.markdown("---")
                docx = create_docx(st.session_state.brief_text, "Сводка (РК)")
                st.download_button("💾 Скачать Сводку (.docx)", docx, "Brief_KZ.docx", use_container_width=True)
            else:
                st.markdown("*Здесь появится результат анализа...*")

# --- МОДУЛЬ 2: ДОКУМЕНТЫ (ДИЗАЙН) ---
def render_doc_analyzer():
    st.markdown("## 📑 Анализ Документов")
    st.caption("Проверка договоров и исков на риски по законодательству РК.")

    with st.container(border=True):
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.markdown("#### Шаг 1. Контекст")
            ctx = st.audio_input("Скажите голосом: Чьи интересы защищаем?", key="doc_ctx")
        
        with col2:
            st.markdown("#### Шаг 2. Файлы")
            files = st.file_uploader("Загрузите документы (PDF, Word)", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)

    if 'doc_res' not in st.session_state: st.session_state.doc_res = None

    if files:
        if st.button("🔍 Начать Анализ", use_container_width=True):
            with st.spinner("⏳ Читаю документы и сверяю с законами РК..."):
                instr = "Общий анализ рисков."
                if ctx: instr = transcribe_audio(ctx)
                
                full_text = ""
                for f in files: full_text += f"\n--- {f.name} ---\n{extract_text_from_file(f)}\n"
                
                model = genai.GenerativeModel('models/gemini-2.0-flash')
                res = model.generate_content(f"Инструкция: {instr}\nДокументы:\n{full_text}\n\nСделай анализ (Сильные стороны / Риски / Вывод) по законам РК.")
                st.session_state.doc_res = res.text

    if st.session_state.doc_res:
        with st.container(border=True):
            st.markdown("### 📝 Результат")
            st.markdown(st.session_state.doc_res)
            st.markdown("---")
            docx = create_docx(st.session_state.doc_res, "Анализ")
            st.download_button("💾 Скачать Отчет (.docx)", docx, "Doc_Analysis.docx", use_container_width=True)

# --- МОДУЛЬ 3: СИМУЛЯТОР (ДИЗАЙН) ---
def render_court_simulator():
    st.markdown("## ⚔️ Судебный Тренажер")
    st.caption("Симуляция заседания с Судьей и Оппонентом. Проверьте свою позицию.")

    if "messages" not in st.session_state: st.session_state.messages = []
    if "sim_active" not in st.session_state: st.session_state.sim_active = False
    if "turn_count" not in st.session_state: st.session_state.turn_count = 0
    if "sim_analysis" not in st.session_state: st.session_state.sim_analysis = None

    # НАСТРОЙКИ
    if not st.session_state.sim_active and not st.session_state.sim_analysis:
        with st.container(border=True):
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("### 1. Ваша роль")
                user_role = st.selectbox("Кого представляете?", ["Истец", "Ответчик"], index=0)
            with col2:
                st.markdown("### 2. Материалы")
                sim_files = st.file_uploader("Загрузите дело", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)
            
            st.markdown("---")
            if st.button("⚖️ Открыть заседание", use_container_width=True, type="primary"):
                if sim_files:
                    ctx = ""
                    for f in sim_files: ctx += extract_text_from_file(f) + "\n"
                    
                    if user_role == "Истец":
                        ai_roles = "Роли ИИ: 1. СУДЬЯ. 2. АДВОКАТ ОТВЕТЧИКА."
                        absent = "Ответчик"
                        starter = "СУДЬЯ"
                    else:
                        ai_roles = "Роли ИИ: 1. СУДЬЯ. 2. АДВОКАТ ИСТЦА."
                        absent = "Истец"
                        starter = "СУДЬЯ"

                    model = genai.GenerativeModel('models/gemini-2.0-flash')
                    prompt = f"""
                    {ai_roles}
                    Пользователь: {user_role}. ВНИМАНИЕ: {absent} отсутствует, говори ТОЛЬКО с пользователем.
                    Контекст: ГПК РК. Материалы: {ctx[:30000]}
                    Начни заседание. Представься как {starter} и задай первый вопрос.
                    """
                    res = model.generate_content(prompt)
                    st.session_state.messages = [{"role": "assistant", "content": res.text}]
                    st.session_state.sim_active = True
                    st.session_state.user_role = user_role
                    st.session_state.ai_roles = ai_roles
                    st.session_state.absent = absent
                    st.rerun()
                else:
                    st.error("Без документов суд не начнется!")

    # ПРОЦЕСС
    if st.session_state.sim_active:
        # Чат
        for msg in st.session_state.messages:
            # Разные иконки
            avatar = "👨‍⚖️" if msg["role"] == "assistant" else "👤"
            with st.chat_message(msg["role"], avatar=avatar):
                st.markdown(msg["content"])
        
        st.write("")
        with st.container(border=True):
            st.markdown("**Ваш ответ суду:**")
            key = f"court_mic_{st.session_state.turn_count}"
            user_audio = st.audio_input("Запись голоса", key=key)

            if user_audio:
                with st.spinner("Суд слушает..."):
                    user_text = transcribe_audio(user_audio)
                    st.session_state.messages.append({"role": "user", "content": f"🗣️ {user_text}"})
                    
                    model = genai.GenerativeModel('models/gemini-2.0-flash')
                    hist = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.messages])
                    
                    prompt = f"""
                    {st.session_state.ai_roles}
                    Правило: ИГНОРИРУЙ отсутствующего {st.session_state.absent}. Говори только с Пользователем.
                    История: {hist}
                    Ответ: "{user_text}"
                    Задача: Кто говорит сейчас (Судья или Оппонент)? Оцени ответ. Задай СЛЕДУЮЩИЙ вопрос.
                    """
                    res = model.generate_content(prompt)
                    st.session_state.messages.append({"role": "assistant", "content": res.text})
                    st.session_state.turn_count += 1
                    st.rerun()

        col_end, _ = st.columns([1, 2])
        with col_end:
            if st.button("🏁 Закончить прения", use_container_width=True):
                with st.spinner("Подготовка разбора..."):
                    model = genai.GenerativeModel('models/gemini-2.0-flash')
                    hist = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.messages])
                    res = model.generate_content(f"Разбор (РК). Роль: {st.session_state.user_role}. История: {hist}\nОтчет: Контекст, Сильные, Слабые, Итог.")
                    st.session_state.sim_analysis = res.text
                    st.session_state.sim_active = False
                    st.rerun()

    # РАЗБОР
    if st.session_state.sim_analysis:
        st.balloons() # Немного анимации для финала
        with st.container(border=True):
            st.markdown("## 📊 Результаты симуляции")
            st.markdown(st.session_state.sim_analysis)
            st.markdown("---")
            docx = create_docx(st.session_state.sim_analysis, "Разбор (РК)")
            st.download_button("💾 Скачать Разбор (.docx)", docx, "Debrief.docx", use_container_width=True, type="primary")
            
            if st.button("🔄 Начать заново"):
                st.session_state.messages = []
                st.session_state.sim_analysis = None
                st.session_state.turn_count = 0
                st.rerun()

# --- МЕНЮ (САЙДБАР) ---
def main():
    st.sidebar.image("https://cdn-icons-png.flaticon.com/512/2600/2600246.png", width=100) # Логотип (можно свой)
    st.sidebar.title("Legal OS")
    st.sidebar.caption("v9.0 • KZ Edition")
    st.sidebar.divider()
    
    mode = st.sidebar.radio("Навигация", ["🕵️‍♂️ Анализатор Дела", "📑 Анализ Документов", "⚔️ Судебный Тренажер"])
    
    st.sidebar.divider()
    st.sidebar.info("Система работает на базе Gemini 2.0 Flash. Данные обрабатываются конфиденциально.")

    if mode == "🕵️‍♂️ Анализатор Дела": render_audio_analyzer()
    elif mode == "📑 Анализ Документов": render_doc_analyzer()
    elif mode == "⚔️ Судебный Тренажер": render_court_simulator()

if __name__ == "__main__":

    main()
